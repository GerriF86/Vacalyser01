# src/utils/export_utils.py
import io
import json
import docx
from fpdf import FPDF
import streamlit as st

def export_json(data: dict):
    """
    Trigger a Streamlit download of data as JSON.
    """
    json_str = json.dumps(data, indent=2)
    st.download_button("Download JSON", json_str, file_name="job_data.json", mime="application/json")

def export_word(data: dict):
    """
    Create a Word document summarizing the job and trigger download.
    """
    doc = docx.Document()
    doc.add_heading(data.get("job_title", "Job Title"), level=0)
    doc.add_paragraph(f"Company: {data.get('company_name', '')}")
    doc.add_paragraph(f"Location: {data.get('city', '')}")
    doc.add_paragraph(f"Role Description:\n{data.get('role_description','')}")
    # Add more fields as needed...
    file_stream = io.BytesIO()
    doc.save(file_stream)
    st.download_button("Download Word", file_stream.getvalue(),
                       file_name="Job_Specification.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def export_pdf(data: dict):
    """
    Create a simple PDF summarizing the job and trigger download.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, txt=data.get("job_title","Job Title"), ln=True)
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt=f"Company: {data.get('company_name','')}", ln=True)
    pdf.cell(0, 10, txt=f"Location: {data.get('city','')}", ln=True)
    pdf.ln(5)
    pdf.multi_cell(0, 10, txt=f"Role Description: {data.get('role_description','')}")
    pdf.ln(5)
    # Convert PDF to binary stream
    pdf_bytes = pdf.output(dest="S").encode('latin-1')
    st.download_button("Download PDF", pdf_bytes,
                       file_name="Job_Specification.pdf",
                       mime="application/pdf")
