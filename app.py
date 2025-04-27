import streamlit as st
import os
import requests
import json
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from src.utils import extraction, preprocessing
from src.session_state import initialize_session_state
from src.services.llm_service import LLMService
import plotly.express as px

# Load environment variables (e.g., for API keys)
load_dotenv()

# Initialize API Keys from environment or Streamlit secrets
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
OLLAMA_API_URL = os.getenv("OLLAMA_API_URL") or st.secrets.get("OLLAMA_API_URL", "http://127.0.0.1:11434")

# Model selection at start
provider = st.radio(
    "Choose Model",
    options=["OpenAI GPT-4o-mini", "Local LLaMA"],
    index=0,
    horizontal=True,
    help="Local uses an Ollama server running a LLaMA model."
)
provider_key = "ollama" if provider.startswith("Local") else "openai"

# Instantiate LLM service based on selection
try:
    llm = LLMService(
        provider=provider_key,
        ollama_api_url=OLLAMA_API_URL or "http://127.0.0.1:11434",
        ollama_model="llama3.2:3b",
        openai_api_key=OPENAI_API_KEY,
        openai_org=st.secrets.get("OPENAI_ORGANIZATION"),
        openai_model="gpt-4o"
    )
except Exception as e:
    st.error(f"Failed to initialize language model: {e}")
    st.stop()

# Helper functions
def apply_base_styling():
    """
    Apply global styling (if any CSS or theme customizations are needed).
    """
    # Example: st.markdown("<style>body{ ... }</style>", unsafe_allow_html=True)
    pass

def show_sidebar_links():
    """Add sidebar title and info (static navigation hints or branding)."""
    st.sidebar.title("Vacalyser Wizard")
    st.sidebar.info("Use the steps below to create a job vacancy.")

def get_from_session_state(key, default=None):
    """Safely get a value from st.session_state."""
    return st.session_state.get(key, default)

def store_in_state(key, value):
    """Store a value in st.session_state."""
    st.session_state[key] = value

def parse_file(uploaded_file, file_name: str = "") -> str:
    """
    Parse contents of an uploaded file (PDF, DOCX, TXT) and return text.
    """
    import os
    ext = os.path.splitext(file_name)[1].lower()
    try:
        if ext == ".pdf":
            from src.utils.extraction import extract_text_from_pdf
            text = extract_text_from_pdf(uploaded_file)
        elif ext == ".docx":
            import docx
            document = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in document.paragraphs])
        elif ext == ".txt":
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        else:
            text = ""
    except Exception as e:
        st.error(f"Error reading file: {e}")
        text = ""
    return text

def match_and_store_keys(raw_text: str, session_keys: list):
    """
    Find lines in raw_text that match known field labels and store values in session_state.
    """
    # Define label prefixes to search for (must include trailing colon if present in text)
    label_map = {
        "job_title": "Job Title:",
        "company_name": "Company Name:",
        "brand_name": "Brand Name:",
        "headquarters_location": "HQ Location:",
        "company_website": "Company Website:",
        "date_of_employment_start": "Date of Employment Start:",
        "job_type": "Job Type:",
        "contract_type": "Contract Type:",
        "job_level": "Job Level:",
        "city": "City (Job Location):",
        "team_structure": "Team Structure:",
        "role_description": "Role Description:",
        "reports_to": "Reports To:",
        "supervises": "Supervises:",
        "role_type": "Role Type:",
        "role_priority_projects": "Priority Projects:",
        "travel_requirements": "Travel Requirements:",
        "work_schedule": "Work Schedule:",
        "role_keywords": "Role Keywords:",
        "decision_making_authority": "Decision Making Authority:",
        "role_performance_metrics": "Role Performance Metrics:",
        "task_list": "Task List:",
        "key_responsibilities": "Key Responsibilities:",
        "technical_tasks": "Technical Tasks:",
        "managerial_tasks": "Managerial Tasks:",
        "administrative_tasks": "Administrative Tasks:",
        "customer_facing_tasks": "Customer-Facing Tasks:",
        "internal_reporting_tasks": "Internal Reporting Tasks:",
        "performance_tasks": "Performance Tasks:",
        "innovation_tasks": "Innovation Tasks:",
        "task_prioritization": "Task Prioritization:",
        "hard_skills": "Hard Skills:",
        "soft_skills": "Soft Skills:",
        "must_have_skills": "Must-Have Skills:",
        "nice_to_have_skills": "Nice-to-Have Skills:",
        "certifications_required": "Certifications Required:",
        "language_requirements": "Language Requirements:",
        "tool_proficiency": "Tool Proficiency:",
        "domain_expertise": "Domain Expertise:",
        "leadership_competencies": "Leadership Competencies:",
        "technical_stack": "Technical Stack:",
        "industry_experience": "Industry Experience:",
        "analytical_skills": "Analytical Skills:",
        "communication_skills": "Communication Skills:",
        "project_management_skills": "Project Management Skills:",
        "soft_requirement_details": "Additional Soft Requirements:",
        "visa_sponsorship": "Visa Sponsorship:",
        "salary_range": "Salary Range:",
        "currency": "Currency:",
        "pay_frequency": "Pay Frequency:",
        "commission_structure": "Commission Structure:",
        "bonus_scheme": "Bonus Scheme:",
        "vacation_days": "Vacation Days:",
        "flexible_hours": "Flexible Hours:",
        "remote_work_policy": "Remote Work Policy:",
        "relocation_assistance": "Relocation Assistance:",
        "childcare_support": "Childcare Support:",
        "recruitment_steps": "Recruitment Steps:",
        "recruitment_timeline": "Recruitment Timeline:",
        "number_of_interviews": "Number of Interviews:",
        "interview_format": "Interview Format:",
        "assessment_tests": "Assessment Tests:",
        "onboarding_process_overview": "Onboarding Process Overview:",
        "recruitment_contact_email": "Recruitment Contact Email:",
        "recruitment_contact_phone": "Recruitment Contact Phone:",
        "application_instructions": "Application Instructions:",
        # Additional metadata fields (these likely won't appear in raw text, but included for completeness)
        "parsed_data_raw": "Parsed Data Raw:",
        "language_of_ad": "Language of Ad:",
        "translation_required": "Translation Required:",
        "employer_branding_elements": "Employer Branding Elements:",
        "desired_publication_channels": "Desired Publication Channels:",
        "internal_job_id": "Internal Job ID:",
        "ad_seniority_tone": "Ad Seniority Tone:",
        "ad_length_preference": "Ad Length Preference:",
        "deadline_urgency": "Deadline Urgency:",
        "company_awards": "Company Awards:",
        "diversity_inclusion_statement": "Diversity & Inclusion Statement:",
        "legal_disclaimers": "Legal Disclaimers:",
        "social_media_links": "Social Media Links:",
        "video_introduction_option": "Video Introduction Option:",
        "comments_internal": "Comments (Internal):"
    }
    for key, label in label_map.items():
        if label in raw_text:
            # Take text after the label up to the next newline
            try:
                value = raw_text.split(label, 1)[1].split("\n", 1)[0].strip()
            except Exception:
                value = raw_text.split(label, 1)[-1].strip()
            if value:
                st.session_state[key] = value

# Keys for all expected session fields (mostly for reference; session state is initialized with these keys)
SESSION_KEYS = [
    "job_title", "company_name", "brand_name", "headquarters_location", "company_website",
    "date_of_employment_start", "job_type", "contract_type", "job_level", "city", "team_structure",
    "role_description", "reports_to", "supervises", "role_type", "role_priority_projects",
    "travel_requirements", "work_schedule", "role_keywords", "decision_making_authority",
    "role_performance_metrics", "task_list", "key_responsibilities", "technical_tasks", "managerial_tasks",
    "administrative_tasks", "customer_facing_tasks", "internal_reporting_tasks", "performance_tasks",
    "innovation_tasks", "task_prioritization", "hard_skills", "soft_skills", "must_have_skills",
    "nice_to_have_skills", "certifications_required", "language_requirements", "tool_proficiency",
    "domain_expertise", "leadership_competencies", "technical_stack", "industry_experience",
    "analytical_skills", "communication_skills", "project_management_skills", "soft_requirement_details",
    "visa_sponsorship", "salary_range", "currency", "pay_frequency", "commission_structure",
    "bonus_scheme", "vacation_days", "flexible_hours", "remote_work_policy", "relocation_assistance",
    "childcare_support", "recruitment_steps", "recruitment_timeline", "number_of_interviews",
    "interview_format", "assessment_tests", "onboarding_process_overview", "recruitment_contact_email",
    "recruitment_contact_phone", "application_instructions", "parsed_data_raw", "language_of_ad",
    "translation_required", "employer_branding_elements", "desired_publication_channels",
    "internal_job_id", "ad_seniority_tone", "ad_length_preference", "deadline_urgency",
    "company_awards", "diversity_inclusion_statement", "legal_disclaimers", "social_media_links",
    "video_introduction_option", "comments_internal",
    # Generated output keys (not initialized at start, will be added if used):
    "target_group_analysis", "generated_job_ad", "generated_interview_prep",
    "generated_email_template", "generated_boolean_query"
]

# Wizard Step 1: Start Discovery
def start_discovery_page():
    """Wizard Step 1: Start Discovery."""
    st.title("Vacalyser")
    st.markdown(
        "**Enhancing hiring workflows** with intelligent suggestions and automations. "
        "We help teams fine-tune job postings and CVs efficiently for better hiring outcomes."
    )
    st.header("1) Start Discovery")
    st.write(
        "Enter a **Job Title**, and optionally provide a link or upload an existing job description file. "
        "We will analyze the provided information to auto-fill relevant fields."
    )
    # Two columns: left for text inputs, right for file upload
    col1, col2 = st.columns(2)
    with col1:
        job_title_val = get_from_session_state("job_title", "")
        job_title = st.text_input("Job Title", value=job_title_val, placeholder="e.g. Senior Data Scientist")
        store_in_state("job_title", job_title)
        default_url = get_from_session_state("input_url", "http://www.")
        input_url = st.text_input("Job Ad / Company URL", value=default_url, placeholder="e.g. https://company.com/careers/job-id")
        store_in_state("input_url", input_url)
    with col2:
        uploaded_file = st.file_uploader("Upload Job Description (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
        if uploaded_file:
            raw_text = parse_file(uploaded_file, uploaded_file.name)
            st.session_state["uploaded_file"] = raw_text
            st.success("File uploaded successfully.")
    # Analyze button
    if st.button("Analyze Sources"):
        raw_text = st.session_state.get("uploaded_file", "")
        # If no file content, try to fetch from URL
        if not raw_text:
            url = input_url.strip()
            if url and url != "http://www.":
                # Prepend scheme if missing
                if not url.lower().startswith("http"):
                    url = "https://" + url
                try:
                    resp = requests.get(url, timeout=10)
                    resp.raise_for_status()
                    soup = BeautifulSoup(resp.text, "html.parser")
                    # Remove scripts and styles
                    for tag in soup(["script", "style"]):
                        tag.decompose()
                    text = soup.get_text(separator="\n")
                    # Keep only non-empty lines
                    lines = [line.strip() for line in text.splitlines() if line.strip()]
                    raw_text = "\n".join(lines)
                except Exception as e:
                    st.error(f"Failed to retrieve content from URL: {e}")
                    raw_text = ""
        if not raw_text:
            st.warning("Please provide a valid URL or upload a file before analysis.")
            return
        # Use LLM to extract structured info as JSON
        # Only include relevant keys (excluding non-input keys)
        extract_keys = [k for k in SESSION_KEYS if k not in ("wizard_step", "input_url", "uploaded_file",
                                                             "target_group_analysis", "generated_job_ad", "generated_interview_prep",
                                                             "generated_email_template", "generated_boolean_query")]
        prompt = (
            f"Extract the following information from the job description below and return it in JSON format with keys {extract_keys}: \n\n"
            f"{raw_text}\n\n"
            f"Output JSON only with the specified keys."
        )
        try:
            llm_response = llm.complete(prompt=prompt)
        except Exception as e:
            st.error(f"Analysis failed: {e}")
            return
        structured_data = {}
        if llm_response:
            try:
                structured_data = json.loads(llm_response)
            except json.JSONDecodeError:
                # If not valid JSON, we'll rely solely on keyword matching
                st.warning("AI parsing did not return valid JSON. Attempting to extract key fields from text.")
        # Update session state with extracted info
        if isinstance(structured_data, dict):
            # Only update keys that exist in session (avoid adding unexpected keys)
            for k, v in structured_data.items():
                if k in st.session_state and v is not None:
                    st.session_state[k] = v
        # Always run keyword matching as a fallback/complement
        match_and_store_keys(raw_text, SESSION_KEYS)
        st.success("Information extracted and fields populated.")
        # Proceed to next step automatically
        st.session_state["wizard_step"] = 2
        st.experimental_rerun()

# Wizard Step 2: Basic Job & Company Info
def render_step_2():
    st.title("Step 2: Basic Job & Company Info")
    st.session_state["job_title"] = st.text_input("Job Title", value=st.session_state["job_title"], placeholder="e.g. Data Scientist")
    st.session_state["company_name"] = st.text_input("Company Name", value=st.session_state["company_name"], placeholder="e.g. TechCorp Ltd.")
    st.session_state["brand_name"] = st.text_input("Brand Name", value=st.session_state["brand_name"], placeholder="If different from company name")
    st.session_state["headquarters_location"] = st.text_input("Headquarters Location", value=st.session_state["headquarters_location"], placeholder="e.g. Berlin, Germany")
    st.session_state["company_website"] = st.text_input("Company Website", value=st.session_state["company_website"], placeholder="e.g. https://techcorp.com")
    st.session_state["date_of_employment_start"] = st.text_input("Preferred/Earliest Start Date", value=st.session_state["date_of_employment_start"], placeholder="e.g. ASAP or 2025-05-01")
    st.session_state["job_type"] = st.selectbox("Job Type", ["Full-Time", "Part-Time", "Internship", "Freelance", "Volunteer", "Other"], index=(["Full-Time", "Part-Time", "Internship", "Freelance", "Volunteer", "Other"].index(st.session_state["job_type"]) if st.session_state["job_type"] else 0))
    st.session_state["contract_type"] = st.selectbox("Contract Type", ["Permanent", "Fixed-Term", "Contract/Freelance", "Other"], index=(["Permanent", "Fixed-Term", "Contract/Freelance", "Other"].index(st.session_state["contract_type"]) if st.session_state["contract_type"] else 0))
    st.session_state["job_level"] = st.selectbox("Job Level", ["Entry-level", "Mid-level", "Senior", "Director", "C-level", "Other"], index=(["Entry-level", "Mid-level", "Senior", "Director", "C-level", "Other"].index(st.session_state["job_level"]) if st.session_state["job_level"] else 0))
    st.session_state["city"] = st.text_input("City (Job Location)", value=st.session_state["city"], placeholder="e.g. Berlin")
    st.session_state["team_structure"] = st.text_area("Team Structure", value=st.session_state["team_structure"], placeholder="Describe how the team is structured...")

# Wizard Step 3: Role Definition
def render_step_3():
    st.title("Step 3: Role Definition")
    st.session_state["role_description"] = st.text_area("Role Description", value=st.session_state["role_description"], placeholder="High-level summary of the role...")
    st.session_state["reports_to"] = st.text_input("Reports To", value=st.session_state["reports_to"], placeholder="Position or name of the manager")
    st.session_state["supervises"] = st.text_area("Supervises", value=st.session_state["supervises"], placeholder="List any roles or teams this role manages")
    st.session_state["role_type"] = st.selectbox("Role Type", ["Individual Contributor", "Team Lead", "Manager", "Director", "Executive", "Other"], index=(["Individual Contributor", "Team Lead", "Manager", "Director", "Executive", "Other"].index(st.session_state["role_type"]) if st.session_state["role_type"] else 0))
    st.session_state["role_priority_projects"] = st.text_area("Priority Projects", value=st.session_state["role_priority_projects"], placeholder="High-priority projects or initiatives for this role")
    st.session_state["travel_requirements"] = st.text_input("Travel Requirements", value=st.session_state["travel_requirements"], placeholder="e.g. Up to 20% travel")
    st.session_state["work_schedule"] = st.text_input("Work Schedule", value=st.session_state["work_schedule"], placeholder="e.g. 9-5 M-F, shift-based, etc.")
    st.session_state["role_keywords"] = st.text_area("Role Keywords", value=st.session_state["role_keywords"], placeholder="Keywords or tags relevant to this role")
    st.session_state["decision_making_authority"] = st.text_input("Decision-Making Authority", value=st.session_state["decision_making_authority"], placeholder="e.g. Budget approvals up to $10k")
    st.session_state["role_performance_metrics"] = st.text_area("Role Performance Metrics", value=st.session_state["role_performance_metrics"], placeholder="Key performance indicators for this role")

# Wizard Step 4: Tasks & Responsibilities
def render_step_4():
    st.title("Step 4: Tasks & Responsibilities")
    st.session_state["task_list"] = st.text_area("General Task List", value=st.session_state["task_list"], placeholder="List of day-to-day tasks (bullets recommended)")
    st.session_state["key_responsibilities"] = st.text_area("Key Responsibilities", value=st.session_state["key_responsibilities"], placeholder="High-level responsibilities of the role")
    st.session_state["technical_tasks"] = st.text_area("Technical Tasks", value=st.session_state["technical_tasks"], placeholder="Technical tasks or duties")
    st.session_state["managerial_tasks"] = st.text_area("Managerial Tasks", value=st.session_state["managerial_tasks"], placeholder="Management or leadership tasks")
    st.session_state["administrative_tasks"] = st.text_area("Administrative Tasks", value=st.session_state["administrative_tasks"], placeholder="Administrative or operational tasks")
    st.session_state["customer_facing_tasks"] = st.text_area("Customer-Facing Tasks", value=st.session_state["customer_facing_tasks"], placeholder="Interactions with clients or customers")
    st.session_state["internal_reporting_tasks"] = st.text_area("Internal Reporting Tasks", value=st.session_state["internal_reporting_tasks"], placeholder="Reporting obligations within the organization")
    st.session_state["performance_tasks"] = st.text_area("Performance-Related Tasks", value=st.session_state["performance_tasks"], placeholder="Tasks related to performance tracking or improvement")
    st.session_state["innovation_tasks"] = st.text_area("Innovation Tasks", value=st.session_state["innovation_tasks"], placeholder="Research & development or innovation tasks")
    st.session_state["task_prioritization"] = st.text_area("Task Prioritization", value=st.session_state["task_prioritization"], placeholder="Information on how tasks are prioritized")

# Wizard Step 5: Skills & Competencies
def render_step_5():
    st.title("Step 5: Skills & Competencies")
    st.session_state["hard_skills"] = st.text_area("Hard Skills", value=st.session_state["hard_skills"], placeholder="Technical or role-specific skills required")
    st.session_state["soft_skills"] = st.text_area("Soft Skills", value=st.session_state["soft_skills"], placeholder="Soft skills (communication, teamwork, etc.)")
    st.session_state["must_have_skills"] = st.text_area("Must-Have Skills", value=st.session_state["must_have_skills"], placeholder="Critical requirements that candidates must have")
    st.session_state["nice_to_have_skills"] = st.text_area("Nice-to-Have Skills", value=st.session_state["nice_to_have_skills"], placeholder="Additional skills that are beneficial but not mandatory")
    st.session_state["certifications_required"] = st.text_area("Certifications Required", value=st.session_state["certifications_required"], placeholder="Any certifications or licenses required for the role")
    st.session_state["language_requirements"] = st.text_area("Language Requirements", value=st.session_state["language_requirements"], placeholder="Language proficiency requirements (e.g. English C1)")
    st.session_state["tool_proficiency"] = st.text_area("Tool Proficiency", value=st.session_state["tool_proficiency"], placeholder="Specific software or tools experience needed")
    st.session_state["domain_expertise"] = st.text_area("Domain Expertise", value=st.session_state["domain_expertise"], placeholder="Industry or domain experience required")
    st.session_state["leadership_competencies"] = st.text_area("Leadership Competencies", value=st.session_state["leadership_competencies"], placeholder="If applicable, leadership or managerial competencies needed")
    st.session_state["technical_stack"] = st.text_area("Technical Stack", value=st.session_state["technical_stack"], placeholder="Applicable tech stack (for technical roles)")
    st.session_state["industry_experience"] = st.text_input("Industry Experience", value=st.session_state["industry_experience"], placeholder="Years of experience in relevant industry")
    st.session_state["analytical_skills"] = st.text_input("Analytical Skills", value=st.session_state["analytical_skills"], placeholder="Analytical or critical-thinking skills")
    st.session_state["communication_skills"] = st.text_input("Communication Skills", value=st.session_state["communication_skills"], placeholder="Communication skills required")
    st.session_state["project_management_skills"] = st.text_input("Project Management Skills", value=st.session_state["project_management_skills"], placeholder="Project management or organizational skills")
    st.session_state["soft_requirement_details"] = st.text_area("Additional Soft Requirements", value=st.session_state["soft_requirement_details"], placeholder="Other personality traits or soft skills desired")
    st.session_state["visa_sponsorship"] = st.selectbox("Visa Sponsorship", ["No", "Yes", "Case-by-Case"], index=(["No", "Yes", "Case-by-Case"].index(st.session_state["visa_sponsorship"]) if st.session_state["visa_sponsorship"] else 0))

# Wizard Step 6: Compensation & Benefits
def render_step_6():
    st.title("Step 6: Compensation & Benefits")
    st.session_state["salary_range"] = st.text_input("Salary Range", value=st.session_state["salary_range"], placeholder="e.g. 50,000 - 60,000")
    st.session_state["currency"] = st.selectbox("Currency", ["EUR", "USD", "GBP", "Other"], index=(["EUR", "USD", "GBP", "Other"].index(st.session_state["currency"]) if st.session_state["currency"] else 0))
    st.session_state["pay_frequency"] = st.selectbox("Pay Frequency", ["Annual", "Monthly", "Bi-weekly", "Weekly", "Other"], index=(["Annual", "Monthly", "Bi-weekly", "Weekly", "Other"].index(st.session_state["pay_frequency"]) if st.session_state["pay_frequency"] else 0))
    st.session_state["commission_structure"] = st.text_input("Commission Structure", value=st.session_state["commission_structure"], placeholder="Details if the role has sales commission")
    st.session_state["bonus_scheme"] = st.text_input("Bonus Scheme", value=st.session_state["bonus_scheme"], placeholder="Details of any performance bonus scheme")
    st.session_state["vacation_days"] = st.text_input("Vacation Days", value=st.session_state["vacation_days"], placeholder="Number of paid vacation days")
    st.session_state["flexible_hours"] = st.selectbox("Flexible Hours", ["No", "Yes", "Partial/Flex Schedule"], index=(["No", "Yes", "Partial/Flex Schedule"].index(st.session_state["flexible_hours"]) if st.session_state["flexible_hours"] else 0))
    st.session_state["remote_work_policy"] = st.selectbox("Remote Work Policy", ["On-site", "Hybrid", "Full Remote", "Other"], index=(["On-site", "Hybrid", "Full Remote", "Other"].index(st.session_state["remote_work_policy"]) if st.session_state["remote_work_policy"] else 0))
    st.session_state["relocation_assistance"] = st.selectbox("Relocation Assistance", ["No", "Yes", "Case-by-Case"], index=(["No", "Yes", "Case-by-Case"].index(st.session_state["relocation_assistance"]) if st.session_state["relocation_assistance"] else 0))
    st.session_state["childcare_support"] = st.selectbox("Childcare Support", ["No", "Yes", "Case-by-Case"], index=(["No", "Yes", "Case-by-Case"].index(st.session_state["childcare_support"]) if st.session_state["childcare_support"] else 0))

# Wizard Step 7: Recruitment Process
def render_step_7():
    st.title("Step 7: Recruitment Process")
    st.session_state["recruitment_steps"] = st.text_area("Recruitment Steps", value=st.session_state["recruitment_steps"], placeholder="e.g. Resume screening, 2 interview rounds, assignment, final interview")
    st.session_state["recruitment_timeline"] = st.text_input("Recruitment Timeline", value=st.session_state["recruitment_timeline"], placeholder="Approximate time from application to final decision")
    st.session_state["number_of_interviews"] = st.text_input("Number of Interviews", value=st.session_state["number_of_interviews"], placeholder="e.g. 2")
    st.session_state["interview_format"] = st.text_input("Interview Format", value=st.session_state["interview_format"], placeholder="e.g. Phone screen, On-site, Video call")
    st.session_state["assessment_tests"] = st.text_area("Assessment Tests", value=st.session_state["assessment_tests"], placeholder="Any tests or assignments (e.g. coding test, case study)")
    st.session_state["onboarding_process_overview"] = st.text_area("Onboarding Process Overview", value=st.session_state["onboarding_process_overview"], placeholder="Brief overview of the onboarding process for the new hire")
    st.session_state["recruitment_contact_email"] = st.text_input("Recruitment Contact Email", value=st.session_state["recruitment_contact_email"], placeholder="Contact email for applicants")
    st.session_state["recruitment_contact_phone"] = st.text_input("Recruitment Contact Phone", value=st.session_state["recruitment_contact_phone"], placeholder="Contact phone for applicants")
    st.session_state["application_instructions"] = st.text_area("Application Instructions", value=st.session_state["application_instructions"], placeholder="How to apply (e.g. apply on website or email resume)")

# Wizard Step 8: Additional Info & Final Summary
def render_step_8():
    st.title("Step 8: Additional Information & Final Review")
    st.subheader("Additional Metadata")
    st.session_state["parsed_data_raw"] = st.text_area("Parsed Data (Raw)", value=st.session_state["parsed_data_raw"], placeholder="(Optional) Raw text or JSON from parsing")
    st.session_state["language_of_ad"] = st.text_input("Language of Ad", value=st.session_state["language_of_ad"], placeholder="e.g. English, German")
    st.session_state["translation_required"] = st.selectbox("Translation Required?", ["No", "Yes"], index=(["No", "Yes"].index(st.session_state["translation_required"]) if st.session_state["translation_required"] else 0))
    st.session_state["employer_branding_elements"] = st.text_area("Employer Branding Elements", value=st.session_state["employer_branding_elements"], placeholder="Company culture or branding points to include in the ad")
    st.session_state["desired_publication_channels"] = st.text_area("Desired Publication Channels", value=st.session_state["desired_publication_channels"], placeholder="Where will this job be advertised? (job boards, social media, etc.)")
    st.session_state["internal_job_id"] = st.text_input("Internal Job ID", value=st.session_state["internal_job_id"], placeholder="Internal reference or job ID")
    st.session_state["ad_seniority_tone"] = st.selectbox("Ad Seniority Tone", ["Casual", "Formal", "Neutral", "Enthusiastic"], index=(["Casual", "Formal", "Neutral", "Enthusiastic"].index(st.session_state["ad_seniority_tone"]) if st.session_state["ad_seniority_tone"] else 0))
    st.session_state["ad_length_preference"] = st.selectbox("Ad Length Preference", ["Short & Concise", "Detailed", "Flexible"], index=(["Short & Concise", "Detailed", "Flexible"].index(st.session_state["ad_length_preference"]) if st.session_state["ad_length_preference"] else 0))
    st.session_state["deadline_urgency"] = st.text_input("Deadline Urgency", value=st.session_state["deadline_urgency"], placeholder="e.g. Position to be filled by end of Q2")
    st.session_state["company_awards"] = st.text_area("Company Awards", value=st.session_state["company_awards"], placeholder="Notable awards or recognitions of the company")
    st.session_state["diversity_inclusion_statement"] = st.text_area("Diversity & Inclusion Statement", value=st.session_state["diversity_inclusion_statement"], placeholder="Equal opportunity employer statement or similar")
    st.session_state["legal_disclaimers"] = st.text_area("Legal Disclaimers", value=st.session_state["legal_disclaimers"], placeholder="Any legal or compliance text to include")
    st.session_state["social_media_links"] = st.text_area("Social Media Links", value=st.session_state["social_media_links"], placeholder="Links to company social media (if relevant for the ad)")
    st.session_state["video_introduction_option"] = st.selectbox("Video Introduction Option?", ["No", "Yes"], index=(["No", "Yes"].index(st.session_state["video_introduction_option"]) if st.session_state["video_introduction_option"] else 0))
    st.session_state["comments_internal"] = st.text_area("Comments (Internal)", value=st.session_state["comments_internal"], placeholder="Any internal comments or notes about this job spec")
    # Final summary preview
    st.subheader("Final Summary (Preview)")
    # Basic info
    st.markdown(f"**Job Title:** {st.session_state['job_title']}")
    st.markdown(f"**Company Name:** {st.session_state['company_name']}")
    st.markdown(f"**Brand Name:** {st.session_state['brand_name']}")
    st.markdown(f"**HQ Location:** {st.session_state['headquarters_location']}")
    st.markdown(f"**Company Website:** {st.session_state['company_website']}")
    st.markdown(f"**Start Date:** {st.session_state['date_of_employment_start']}")
    st.markdown(f"**Job Type:** {st.session_state['job_type']}")
    st.markdown(f"**Contract Type:** {st.session_state['contract_type']}")
    st.markdown(f"**Job Level:** {st.session_state['job_level']}")
    st.markdown(f"**Job Location (City):** {st.session_state['city']}")
    st.markdown(f"**Team Structure:** {st.session_state['team_structure']}")
    st.markdown("---")
    # Role definition
    st.markdown(f"**Role Description:** {st.session_state['role_description']}")
    st.markdown(f"**Reports To:** {st.session_state['reports_to']}")
    st.markdown(f"**Supervises:** {st.session_state['supervises']}")
    st.markdown(f"**Role Type:** {st.session_state['role_type']}")
    st.markdown(f"**Priority Projects:** {st.session_state['role_priority_projects']}")
    st.markdown(f"**Travel Requirements:** {st.session_state['travel_requirements']}")
    st.markdown(f"**Work Schedule:** {st.session_state['work_schedule']}")
    st.markdown(f"**Role Keywords:** {st.session_state['role_keywords']}")
    st.markdown(f"**Decision Making Authority:** {st.session_state['decision_making_authority']}")
    st.markdown(f"**Performance Metrics:** {st.session_state['role_performance_metrics']}")
    st.markdown("---")
    # Tasks & Responsibilities
    st.markdown(f"**General Task List:** {st.session_state['task_list']}")
    st.markdown(f"**Key Responsibilities:** {st.session_state['key_responsibilities']}")
    st.markdown(f"**Technical Tasks:** {st.session_state['technical_tasks']}")
    st.markdown(f"**Managerial Tasks:** {st.session_state['managerial_tasks']}")
    st.markdown(f"**Administrative Tasks:** {st.session_state['administrative_tasks']}")
    st.markdown(f"**Customer-Facing Tasks:** {st.session_state['customer_facing_tasks']}")
    st.markdown(f"**Internal Reporting Tasks:** {st.session_state['internal_reporting_tasks']}")
    st.markdown(f"**Performance-Related Tasks:** {st.session_state['performance_tasks']}")
    st.markdown(f"**Innovation Tasks:** {st.session_state['innovation_tasks']}")
    st.markdown(f"**Task Prioritization:** {st.session_state['task_prioritization']}")
    st.markdown("---")
    # Skills & Competencies
    st.markdown(f"**Hard Skills:** {st.session_state['hard_skills']}")
    st.markdown(f"**Soft Skills:** {st.session_state['soft_skills']}")
    st.markdown(f"**Must-Have Skills:** {st.session_state['must_have_skills']}")
    st.markdown(f"**Nice-to-Have Skills:** {st.session_state['nice_to_have_skills']}")
    st.markdown(f"**Certifications Required:** {st.session_state['certifications_required']}")
    st.markdown(f"**Language Requirements:** {st.session_state['language_requirements']}")
    st.markdown(f"**Tool Proficiency:** {st.session_state['tool_proficiency']}")
    st.markdown(f"**Domain Expertise:** {st.session_state['domain_expertise']}")
    st.markdown(f"**Leadership Competencies:** {st.session_state['leadership_competencies']}")
    st.markdown(f"**Technical Stack:** {st.session_state['technical_stack']}")
    st.markdown(f"**Industry Experience:** {st.session_state['industry_experience']}")
    st.markdown(f"**Analytical Skills:** {st.session_state['analytical_skills']}")
    st.markdown(f"**Communication Skills:** {st.session_state['communication_skills']}")
    st.markdown(f"**Project Management Skills:** {st.session_state['project_management_skills']}")
    st.markdown(f"**Additional Soft Requirements:** {st.session_state['soft_requirement_details']}")
    st.markdown(f"**Visa Sponsorship:** {st.session_state['visa_sponsorship']}")
    st.markdown("---")
    # Compensation & Benefits
    st.markdown(f"**Salary Range:** {st.session_state['salary_range']} {st.session_state['currency']}")
    st.markdown(f"**Pay Frequency:** {st.session_state['pay_frequency']}")
    st.markdown(f"**Commission Structure:** {st.session_state['commission_structure']}")
    st.markdown(f"**Bonus Scheme:** {st.session_state['bonus_scheme']}")
    st.markdown(f"**Vacation Days:** {st.session_state['vacation_days']}")
    st.markdown(f"**Flexible Hours:** {st.session_state['flexible_hours']}")
    st.markdown(f"**Remote Work Policy:** {st.session_state['remote_work_policy']}")
    st.markdown(f"**Relocation Assistance:** {st.session_state['relocation_assistance']}")
    st.markdown(f"**Childcare Support:** {st.session_state['childcare_support']}")
    st.markdown("---")
    # Recruitment Process
    st.markdown(f"**Recruitment Steps:** {st.session_state['recruitment_steps']}")
    st.markdown(f"**Recruitment Timeline:** {st.session_state['recruitment_timeline']}")
    st.markdown(f"**Number of Interviews:** {st.session_state['number_of_interviews']}")
    st.markdown(f"**Interview Format:** {st.session_state['interview_format']}")
    st.markdown(f"**Assessment Tests:** {st.session_state['assessment_tests']}")
    st.markdown(f"**Onboarding Process Overview:** {st.session_state['onboarding_process_overview']}")
    st.markdown(f"**Contact Email:** {st.session_state['recruitment_contact_email']}")
    st.markdown(f"**Contact Phone:** {st.session_state['recruitment_contact_phone']}")
    st.markdown(f"**Application Instructions:** {st.session_state['application_instructions']}")
    st.markdown("---")
    # Additional Metadata
    st.markdown(f"**Parsed Data (Raw):** {st.session_state['parsed_data_raw']}")
    st.markdown(f"**Language of Ad:** {st.session_state['language_of_ad']}")
    st.markdown(f"**Translation Required:** {st.session_state['translation_required']}")
    st.markdown(f"**Employer Branding Elements:** {st.session_state['employer_branding_elements']}")
    st.markdown(f"**Desired Publication Channels:** {st.session_state['desired_publication_channels']}")
    st.markdown(f"**Internal Job ID:** {st.session_state['internal_job_id']}")
    st.markdown(f"**Ad Seniority Tone:** {st.session_state['ad_seniority_tone']}")
    st.markdown(f"**Ad Length Preference:** {st.session_state['ad_length_preference']}")
    st.markdown(f"**Deadline Urgency:** {st.session_state['deadline_urgency']}")
    st.markdown(f"**Company Awards:** {st.session_state['company_awards']}")
    st.markdown(f"**Diversity & Inclusion Statement:** {st.session_state['diversity_inclusion_statement']}")
    st.markdown(f"**Legal Disclaimers:** {st.session_state['legal_disclaimers']}")
    st.markdown(f"**Social Media Links:** {st.session_state['social_media_links']}")
    st.markdown(f"**Video Introduction Option:** {st.session_state['video_introduction_option']}")
    st.markdown(f"**Internal Comments:** {st.session_state['comments_internal']}")
    st.info("Review the information above. You can go back to previous steps to make changes if needed, or generate the final outputs now.")
    # AI-Generated Outputs
    st.subheader("AI-Generated Outputs")
    tabs = st.tabs(["Target Group Analysis", "Job Advertisement", "Interview Prep"])
    # Target Group Analysis generation
    with tabs[0]:
        if st.button("Generate Target Group Analysis"):
            # Prompt for target group analysis
            info = f"Job Title: {st.session_state['job_title']}\n"
            if st.session_state["company_name"]:
                info += f"Company: {st.session_state['company_name']}\n"
            if st.session_state["job_level"]:
                info += f"Level: {st.session_state['job_level']}\n"
            # Include key skills
            key_skills = st.session_state["must_have_skills"] or st.session_state["hard_skills"]
if st.session_state.get("generated_prep"):
    st.subheader("Generated Interview Prep Guide")
    prep_text = st.session_state["generated_prep"]
    st.text_area("Interview Prep", prep_text, height=200)
    st.download_button(
        "Download Prep Guide (TXT)",
        prep_text,
        file_name="InterviewPrep.txt"
    )

# Attempt to import PDF generator
try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

def text_to_pdf_bytes(text: str) -> bytes:
    """Convert text to PDF bytes (simple layout). Returns None if FPDF not available."""
    if not FPDF:
        return None
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Multi-cell to handle multi-line text
    pdf.multi_cell(0, 10, text)
    return pdf.output(dest='S').encode('latin-1', errors='ignore')

# Load environment variables (for API keys, etc.)
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
OLLAMA_API_URL = os.getenv("OLLAMA_API_URL") or st.secrets.get("OLLAMA_API_URL", "http://127.0.0.1:11434")

# Model selection
provider = st.radio(
    "Choose Model",
    options=["OpenAI GPT-4o-mini", "Local LLaMA"],
    index=0,
    horizontal=True,
    help="Local uses an Ollama server running a LLaMA model."
)
provider_key = "ollama" if provider.startswith("Local") else "openai"
# Initialize LLM service
try:
    llm = LLMService(
        provider=provider_key,
        ollama_api_url=OLLAMA_API_URL or "http://127.0.0.1:11434",
        ollama_model="llama3.2:3b",
        openai_api_key=OPENAI_API_KEY,
        openai_org=st.secrets.get("OPENAI_ORGANIZATION"),
        openai_model="gpt-4o"
    )
except Exception as e:
    st.error(f"Failed to initialize LLM: {e}")
    st.stop()

def parse_file(uploaded_file) -> str:
    """Read uploaded file (pdf/docx/txt) and return text."""
    import os
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    try:
        if ext == ".pdf":
            from src.utils.extraction import extract_text_from_pdf
            text = extract_text_from_pdf(uploaded_file)
        elif ext == ".docx":
            import docx
            doc = docx.Document(uploaded_file)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif ext == ".txt":
            text = uploaded_file.read().decode("utf-8", errors="ignore")
        else:
            text = ""
    except Exception as err:
        st.error(f"Error reading file: {err}")
        text = ""
    return text

def match_and_store_keys(raw_text: str):
    """Find known labels in raw_text and store their values in session_state."""
    label_map = {
        "job_title": "Job Title:", "company_name": "Company Name:", "brand_name": "Brand Name:",
        "headquarters_location": "HQ Location:", "company_website": "Company Website:",
        "date_of_employment_start": "Date of Employment Start:", "job_type": "Job Type:",
        "contract_type": "Contract Type:", "job_level": "Job Level:", "city": "City",
        "team_structure": "Team Structure:", "role_description": "Role Description:",
        "reports_to": "Reports To:", "supervises": "Supervises:", "role_type": "Role Type:",
        "role_priority_projects": "Priority Projects:", "travel_requirements": "Travel Requirements:",
        "work_schedule": "Work Schedule:", "role_keywords": "Role Keywords:",
        "decision_making_authority": "Decision Making Authority:", "role_performance_metrics": "Role Performance Metrics:",
        "task_list": "Task List:", "key_responsibilities": "Key Responsibilities:",
        "technical_tasks": "Technical Tasks:", "managerial_tasks": "Managerial Tasks:",
        "administrative_tasks": "Administrative Tasks:", "customer_facing_tasks": "Customer-Facing Tasks:",
        "internal_reporting_tasks": "Internal Reporting Tasks:", "performance_tasks": "Performance Tasks:",
        "innovation_tasks": "Innovation Tasks:", "task_prioritization": "Task Prioritization:",
        "hard_skills": "Hard Skills:", "soft_skills": "Soft Skills:",
        "must_have_skills": "Must-Have Skills:", "nice_to_have_skills": "Nice-to-Have Skills:",
        "certifications_required": "Certifications Required:", "language_requirements": "Language Requirements:",
        "tool_proficiency": "Tool Proficiency:", "domain_expertise": "Domain Expertise:",
        "leadership_competencies": "Leadership Competencies:", "technical_stack": "Technical Stack:",
        "industry_experience": "Industry Experience:", "analytical_skills": "Analytical Skills:",
        "communication_skills": "Communication Skills:", "project_management_skills": "Project Management Skills:",
        "soft_requirement_details": "Additional Soft Requirements:", "visa_sponsorship": "Visa Sponsorship:",
        "salary_range": "Salary Range:", "currency": "Currency:", "pay_frequency": "Pay Frequency:",
        "commission_structure": "Commission Structure:", "bonus_scheme": "Bonus Scheme:",
        "vacation_days": "Vacation Days:", "flexible_hours": "Flexible Hours:",
        "remote_work_policy": "Remote Work Policy:", "relocation_assistance": "Relocation Assistance:",
        "childcare_support": "Childcare Support:", "recruitment_steps": "Recruitment Steps:",
        "recruitment_timeline": "Recruitment Timeline:", "number_of_interviews": "Number of Interviews:",
        "interview_format": "Interview Format:", "assessment_tests": "Assessment Tests:",
        "onboarding_process_overview": "Onboarding Process Overview:", "recruitment_contact_email": "Recruitment Contact Email:",
        "recruitment_contact_phone": "Recruitment Contact Phone:", "application_instructions": "Application Instructions:"
    }
    for key, label in label_map.items():
        if label in raw_text:
            try:
                value = raw_text.split(label, 1)[1].split("\n", 1)[0].strip()
            except Exception:
                value = raw_text.split(label, 1)[-1].strip()
            if value:
                st.session_state[key] = value

def start_discovery_page():
    """Step 1: Start Discovery."""
    st.title("Vacalyser")
    st.markdown("**Enhancing hiring workflows** with intelligent suggestions and automations. We help teams fine-tune job postings and CVs for better hiring outcomes.")
    st.header("1) Start Discovery")
    st.write("Enter a **Job Title**, and optionally a link or upload a file for an existing job description. We'll analyze the input to pre-fill fields.")
    col1, col2 = st.columns(2)
    with col1:
        st.text_input("Job Title", placeholder="e.g. Senior Data Scientist", key="job_title")
        st.text_input("Job Ad / Company URL", placeholder="e.g. https://company.com/careers/job-id", key="input_url")
    with col2:
        uploaded_file = st.file_uploader("Upload Job Description (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
        if uploaded_file:
            raw_text = parse_file(uploaded_file)
            st.session_state["uploaded_file"] = raw_text
            st.success("File uploaded successfully.")
    if st.button("Analyze Sources"):
        raw_text = st.session_state.get("uploaded_file", "")
        # If no file text, try URL
        if not raw_text:
            url = st.session_state.get("input_url", "").strip()
            if url and url != "http://www.":
                if not url.lower().startswith("http"):
                    url = "https://" + url
                try:
                    resp = requests.get(url, timeout=10)
                    resp.raise_for_status()
                    soup = BeautifulSoup(resp.text, "html.parser")
                    for tag in soup(["script", "style"]):
                        tag.decompose()
                    text = soup.get_text(separator="\n")
                    lines = [line.strip() for line in text.splitlines() if line.strip()]
                    raw_text = "\n".join(lines)
                except Exception as e:
                    st.error(f"Failed to fetch URL content: {e}")
                    raw_text = ""
        if not raw_text:
            st.warning("Please enter a valid URL or upload a file before analysis.")
            return
        # LLM extraction prompt
        fields_to_extract = [k for k in st.session_state.keys() if k not in ("wizard_step", "input_url", "uploaded_file", 
                                                                            "target_group_analysis", "generated_job_ad", 
                                                                            "generated_interview_prep", "generated_email_template", 
                                                                            "generated_boolean_query")]
        prompt = (
            f"Extract the following information from the job description below and return as JSON with these keys: {fields_to_extract}.\n\n"
            f"{raw_text}\n\nOutput JSON:"
        )
        try:
            llm_response = llm.complete(prompt=prompt)
        except Exception as e:
            st.error(f"Analysis failed: {e}")
            return
        if llm_response:
            try:
                extracted = json.loads(llm_response)
            except json.JSONDecodeError:
                extracted = {}
            if isinstance(extracted, dict):
                for k, v in extracted.items():
                    if k in st.session_state and v is not None:
                        st.session_state[k] = v
        # Fallback keyword matching
        match_and_store_keys(raw_text)
        st.success("Information extracted and fields populated.")
        st.session_state["wizard_step"] = 2
        st.experimental_rerun()

def render_step_2():
    st.title("Step 2: Basic Job & Company Info")
    st.text_input("Job Title", key="job_title")
    st.text_input("Company Name", key="company_name")
    st.text_input("Brand Name", key="brand_name", placeholder="(if different from company)")
    st.text_input("Headquarters Location", key="headquarters_location", placeholder="e.g. City, Country")
    st.text_input("Company Website", key="company_website", placeholder="e.g. https://example.com")
    st.text_input("Preferred/Earliest Start Date", key="date_of_employment_start", placeholder="e.g. ASAP or 2025-05-01")
    st.selectbox("Job Type", ["Full-Time", "Part-Time", "Internship", "Freelance", "Volunteer", "Other"], key="job_type")
    st.selectbox("Contract Type", ["Permanent", "Fixed-Term", "Contract/Freelance", "Other"], key="contract_type")
    st.selectbox("Job Level", ["Entry-level", "Mid-level", "Senior", "Director", "C-level", "Other"], key="job_level")
    st.text_input("City (Job Location)", key="city")

    st.text_area("Team Structure", key="team_structure", placeholder="Describe how the team is structured...")

def render_step_3():
    st.title("Step 3: Role Definition")
    st.text_area("Role Description", key="role_description", placeholder="High-level role summary...")
    st.text_input("Reports To", key="reports_to", placeholder="Manager's position or name")
    st.text_area("Supervises", key="supervises", placeholder="Roles or team this position manages")
    st.selectbox("Role Type", ["Individual Contributor", "Team Lead", "Manager", "Director", "Executive", "Other"], key="role_type")
    st.text_area("Priority Projects", key="role_priority_projects", placeholder="High-priority projects or goals for this role")
    st.text_input("Travel Requirements", key="travel_requirements", placeholder="e.g. Up to 20% travel")
    st.text_input("Work Schedule", key="work_schedule", placeholder="e.g. 9-5 M-F, shift-based, etc.")
    st.text_area("Role Keywords", key="role_keywords", placeholder="Keywords for this role (for search/tags)")
    st.text_input("Decision-Making Authority", key="decision_making_authority", placeholder="e.g. Budget approvals up to X")
    st.text_area("Role Performance Metrics", key="role_performance_metrics", placeholder="Key performance indicators for this role")

def render_step_4():
    st.title("Step 4: Tasks & Responsibilities")
    st.text_area("General Task List", key="task_list", placeholder="Day-to-day tasks (bullet points)")
    st.text_area("Key Responsibilities", key="key_responsibilities", placeholder="Major responsibilities of the role")
    st.text_area("Technical Tasks", key="technical_tasks", placeholder="Technical duties")
    st.text_area("Managerial Tasks", key="managerial_tasks", placeholder="Management duties")
    st.text_area("Administrative Tasks", key="administrative_tasks", placeholder="Administrative duties")
    st.text_area("Customer-Facing Tasks", key="customer_facing_tasks", placeholder="Client/customer interaction tasks")
    st.text_area("Internal Reporting Tasks", key="internal_reporting_tasks", placeholder="Internal reporting duties")
    st.text_area("Performance-Related Tasks", key="performance_tasks", placeholder="Tasks tied to performance metrics")
    st.text_area("Innovation Tasks", key="innovation_tasks", placeholder="R&D or innovation-related tasks")
    st.text_area("Task Prioritization", key="task_prioritization", placeholder="How tasks are prioritized")

def render_step_5():
    st.title("Step 5: Skills & Competencies")
    st.text_area("Hard Skills", key="hard_skills", placeholder="Technical or domain-specific skills required")
    st.text_area("Soft Skills", key="soft_skills", placeholder="Soft skills (communication, teamwork, etc.)")
    st.text_area("Must-Have Skills", key="must_have_skills", placeholder="Critical required skills/qualifications")
    st.text_area("Nice-to-Have Skills", key="nice_to_have_skills", placeholder="Additional beneficial skills")
    st.text_area("Certifications Required", key="certifications_required", placeholder="e.g. PMP, CPA, etc.")
    st.text_area("Language Requirements", key="language_requirements", placeholder="e.g. English C1, German B2")
    st.text_area("Tool Proficiency", key="tool_proficiency", placeholder="Software/tools experience needed (e.g. Excel, AWS)")
    st.text_area("Domain Expertise", key="domain_expertise", placeholder="Industry or domain experience required")
    st.text_area("Leadership Competencies", key="leadership_competencies", placeholder="Leadership/management competencies (if applicable)")
    st.text_area("Technical Stack", key="technical_stack", placeholder="Relevant tech stack (for technical roles)")
    st.text_input("Industry Experience", key="industry_experience", placeholder="Years of experience in relevant industry")
    st.text_input("Analytical Skills", key="analytical_skills", placeholder="Analytical or problem-solving skills")
    st.text_input("Communication Skills", key="communication_skills", placeholder="Communication skills required")
    st.text_input("Project Management Skills", key="project_management_skills", placeholder="Project management skills if required")
    st.text_area("Additional Soft Requirements", key="soft_requirement_details", placeholder="Other soft skills or traits")
    st.selectbox("Visa Sponsorship", ["No", "Yes", "Case-by-Case"], key="visa_sponsorship")

def render_step_6():
    st.title("Step 6: Compensation & Benefits")
    st.text_input("Salary Range", key="salary_range", placeholder="e.g. 50,000 - 60,000")
    st.selectbox("Currency", ["EUR", "USD", "GBP", "Other"], key="currency")
    st.selectbox("Pay Frequency", ["Annual", "Monthly", "Bi-weekly", "Weekly", "Other"], key="pay_frequency")
    st.text_input("Commission Structure", key="commission_structure", placeholder="If applicable, details of sales commission")
    st.text_input("Bonus Scheme", key="bonus_scheme", placeholder="If applicable, details of bonus scheme")
    st.text_input("Vacation Days", key="vacation_days", placeholder="Number of paid vacation days")
    st.selectbox("Flexible Hours", ["No", "Yes", "Partial/Flex Schedule"], key="flexible_hours")
    st.selectbox("Remote Work Policy", ["On-site", "Hybrid", "Full Remote", "Other"], key="remote_work_policy")
    st.selectbox("Relocation Assistance", ["No", "Yes", "Case-by-Case"], key="relocation_assistance")
    st.selectbox("Childcare Support", ["No", "Yes", "Case-by-Case"], key="childcare_support")

def render_step_7():
    st.title("Step 7: Recruitment Process")
    st.text_area("Recruitment Steps", key="recruitment_steps", placeholder="e.g. CV screening, 2 interviews, assignment, final interview")
    st.text_input("Recruitment Timeline", key="recruitment_timeline", placeholder="e.g. ~4 weeks from application to offer")
    st.text_input("Number of Interviews", key="number_of_interviews", placeholder="e.g. 2")
    st.text_input("Interview Format", key="interview_format", placeholder="e.g. Phone screen, then on-site")
    st.text_area("Assessment Tests", key="assessment_tests", placeholder="Any tests or assignments (e.g. coding test)")
    st.text_area("Onboarding Process Overview", key="onboarding_process_overview", placeholder="Brief overview of the onboarding process")
    st.text_input("Recruitment Contact Email", key="recruitment_contact_email", placeholder="Contact email for applicants")
    st.text_input("Recruitment Contact Phone", key="recruitment_contact_phone", placeholder="Contact phone for applicants")
    st.text_area("Application Instructions", key="application_instructions", placeholder="How to apply (e.g. apply on website or send CV)")

def render_step_8():
    st.title("Step 8: Additional Information & Final Review")
    st.subheader("Additional Metadata")
    st.text_area("Parsed Data (Raw)", key="parsed_data_raw", placeholder="(Optional) Raw text or notes")
    st.text_input("Language of Ad", key="language_of_ad", placeholder="e.g. English")
    st.selectbox("Translation Required?", ["No", "Yes"], key="translation_required")
    st.text_area("Employer Branding Elements", key="employer_branding_elements", placeholder="Company culture or branding points to highlight")
    st.text_area("Desired Publication Channels", key="desired_publication_channels", placeholder="Where to publish this job (internal, LinkedIn, job boards, etc.)")
    st.text_input("Internal Job ID", key="internal_job_id", placeholder="Internal reference ID for this job")
    st.selectbox("Ad Seniority Tone", ["Casual", "Formal", "Neutral", "Enthusiastic"], key="ad_seniority_tone")
    st.selectbox("Ad Length Preference", ["Short & Concise", "Detailed", "Flexible"], key="ad_length_preference")
    st.text_input("Deadline Urgency", key="deadline_urgency", placeholder="e.g. Position to be filled by end of Q2")
    st.text_area("Company Awards", key="company_awards", placeholder="Notable company awards or achievements")
    st.text_area("Diversity & Inclusion Statement", key="diversity_inclusion_statement", placeholder="Equal opportunity employer statement, etc.")
    st.text_area("Legal Disclaimers", key="legal_disclaimers", placeholder="Any legal disclaimers to include")
    st.text_area("Social Media Links", key="social_media_links", placeholder="Company social media links (for the ad)")
    st.selectbox("Video Introduction Option?", ["No", "Yes"], key="video_introduction_option")
    st.text_area("Comments (Internal)", key="comments_internal", placeholder="Internal comments or notes")
    # Final summary preview
    st.subheader("Final Summary (Preview)")
    # Basic job info
    st.markdown(f"**Job Title:** {st.session_state['job_title']}")
    st.markdown(f"**Company Name:** {st.session_state['company_name']}")
    st.markdown(f"**Brand Name:** {st.session_state['brand_name']}")
    st.markdown(f"**HQ Location:** {st.session_state['headquarters_location']}")
    st.markdown(f"**Company Website:** {st.session_state['company_website']}")
    st.markdown(f"**Start Date:** {st.session_state['date_of_employment_start']}")
    st.markdown(f"**Job Type:** {st.session_state['job_type']}")
    st.markdown(f"**Contract Type:** {st.session_state['contract_type']}")
    st.markdown(f"**Job Level:** {st.session_state['job_level']}")
    st.markdown(f"**Job Location (City):** {st.session_state['city']}")
    st.markdown(f"**Team Structure:** {st.session_state['team_structure']}")
    st.markdown("---")
    # Role info
    st.markdown(f"**Role Description:** {st.session_state['role_description']}")
    st.markdown(f"**Reports To:** {st.session_state['reports_to']}")
    st.markdown(f"**Supervises:** {st.session_state['supervises']}")
    st.markdown(f"**Role Type:** {st.session_state['role_type']}")
    st.markdown(f"**Priority Projects:** {st.session_state['role_priority_projects']}")
    st.markdown(f"**Travel Requirements:** {st.session_state['travel_requirements']}")
    st.markdown(f"**Work Schedule:** {st.session_state['work_schedule']}")
    st.markdown(f"**Role Keywords:** {st.session_state['role_keywords']}")
    st.markdown(f"**Decision Making Authority:** {st.session_state['decision_making_authority']}")
    st.markdown(f"**Performance Metrics:** {st.session_state['role_performance_metrics']}")
    st.markdown("---")
    # Tasks & responsibilities
    st.markdown(f"**General Task List:** {st.session_state['task_list']}")
    st.markdown(f"**Key Responsibilities:** {st.session_state['key_responsibilities']}")
    st.markdown(f"**Technical Tasks:** {st.session_state['technical_tasks']}")
    st.markdown(f"**Managerial Tasks:** {st.session_state['managerial_tasks']}")
    st.markdown(f"**Administrative Tasks:** {st.session_state['administrative_tasks']}")
    st.markdown(f"**Customer-Facing Tasks:** {st.session_state['customer_facing_tasks']}")
    st.markdown(f"**Internal Reporting Tasks:** {st.session_state['internal_reporting_tasks']}")
    st.markdown(f"**Performance-Related Tasks:** {st.session_state['performance_tasks']}")
    st.markdown(f"**Innovation Tasks:** {st.session_state['innovation_tasks']}")
    st.markdown(f"**Task Prioritization:** {st.session_state['task_prioritization']}")
    st.markdown("---")
    # Skills & competencies
    st.markdown(f"**Hard Skills:** {st.session_state['hard_skills']}")
    st.markdown(f"**Soft Skills:** {st.session_state['soft_skills']}")
    st.markdown(f"**Must-Have Skills:** {st.session_state['must_have_skills']}")
    st.markdown(f"**Nice-to-Have Skills:** {st.session_state['nice_to_have_skills']}")
    st.markdown(f"**Certifications Required:** {st.session_state['certifications_required']}")
    st.markdown(f"**Language Requirements:** {st.session_state['language_requirements']}")
    st.markdown(f"**Tool Proficiency:** {st.session_state['tool_proficiency']}")
    st.markdown(f"**Domain Expertise:** {st.session_state['domain_expertise']}")
    st.markdown(f"**Leadership Competencies:** {st.session_state['leadership_competencies']}")
    st.markdown(f"**Technical Stack:** {st.session_state['technical_stack']}")
    st.markdown(f"**Industry Experience:** {st.session_state['industry_experience']}")
    st.markdown(f"**Analytical Skills:** {st.session_state['analytical_skills']}")
    st.markdown(f"**Communication Skills:** {st.session_state['communication_skills']}")
    st.markdown(f"**Project Management Skills:** {st.session_state['project_management_skills']}")
    st.markdown(f"**Additional Soft Requirements:** {st.session_state['soft_requirement_details']}")
    st.markdown(f"**Visa Sponsorship:** {st.session_state['visa_sponsorship']}")
    st.markdown("---")
    # Compensation & benefits
    st.markdown(f"**Salary Range:** {st.session_state['salary_range']} {st.session_state['currency']}")
    st.markdown(f"**Pay Frequency:** {st.session_state['pay_frequency']}")
    st.markdown(f"**Commission Structure:** {st.session_state['commission_structure']}")
    st.markdown(f"**Bonus Scheme:** {st.session_state['bonus_scheme']}")
    st.markdown(f"**Vacation Days:** {st.session_state['vacation_days']}")
    st.markdown(f"**Flexible Hours:** {st.session_state['flexible_hours']}")
    st.markdown(f"**Remote Work Policy:** {st.session_state['remote_work_policy']}")
    st.markdown(f"**Relocation Assistance:** {st.session_state['relocation_assistance']}")
    st.markdown(f"**Childcare Support:** {st.session_state['childcare_support']}")
    st.markdown("---")
    # Recruitment process
    st.markdown(f"**Recruitment Steps:** {st.session_state['recruitment_steps']}")
    st.markdown(f"**Recruitment Timeline:** {st.session_state['recruitment_timeline']}")
    st.markdown(f"**Number of Interviews:** {st.session_state['number_of_interviews']}")
    st.markdown(f"**Interview Format:** {st.session_state['interview_format']}")
    st.markdown(f"**Assessment Tests:** {st.session_state['assessment_tests']}")
    st.markdown(f"**Onboarding Process Overview:** {st.session_state['onboarding_process_overview']}")
    st.markdown(f"**Contact Email:** {st.session_state['recruitment_contact_email']}")
    st.markdown(f"**Contact Phone:** {st.session_state['recruitment_contact_phone']}")
    st.markdown(f"**Application Instructions:** {st.session_state['application_instructions']}")
    st.markdown("---")
    # Additional metadata
    st.markdown(f"**Parsed Data (Raw):** {st.session_state['parsed_data_raw']}")
    st.markdown(f"**Language of Ad:** {st.session_state['language_of_ad']}")
    st.markdown(f"**Translation Required:** {st.session_state['translation_required']}")
    st.markdown(f"**Employer Branding Elements:** {st.session_state['employer_branding_elements']}")
    st.markdown(f"**Desired Publication Channels:** {st.session_state['desired_publication_channels']}")
    st.markdown(f"**Internal Job ID:** {st.session_state['internal_job_id']}")
    st.markdown(f"**Ad Seniority Tone:** {st.session_state['ad_seniority_tone']}")
    st.markdown(f"**Ad Length Preference:** {st.session_state['ad_length_preference']}")
    st.markdown(f"**Deadline Urgency:** {st.session_state['deadline_urgency']}")
    st.markdown(f"**Company Awards:** {st.session_state['company_awards']}")
    st.markdown(f"**Diversity & Inclusion Statement:** {st.session_state['diversity_inclusion_statement']}")
    st.markdown(f"**Legal Disclaimers:** {st.session_state['legal_disclaimers']}")
    st.markdown(f"**Social Media Links:** {st.session_state['social_media_links']}")
    st.markdown(f"**Video Introduction Option:** {st.session_state['video_introduction_option']}")
    st.markdown(f"**Comments (Internal):** {st.session_state['comments_internal']}")
    st.info("Review the information above. You can go back to edit, or generate the final outputs now.")
    # AI-Generated Outputs
    st.subheader("AI-Generated Outputs")
    tab1, tab2, tab3 = st.tabs(["Target Group Analysis", "Job Advertisement", "Interview Prep"])
    with tab1:
        if st.button("Generate Target Group Analysis"):
            # Prepare prompt for target group analysis
            job_title = st.session_state['job_title']
            company = st.session_state['company_name']
            level = st.session_state['job_level']
            must_skills = st.session_state['must_have_skills']
            # Compose prompt
            prompt_tga = (f"Position: {job_title}\n"
                          + (f"Company: {company}\n" if company else "")
                          + (f"Level: {level}\n" if level else "")
                          + (f"Required Skills: {must_skills}\n" if must_skills else "")
                          + "Provide an analysis of the target candidate group for this position, including the ideal candidate profile (background, experience, skills, motivations) and how to attract them.")
            try:
                result = llm.complete(prompt=prompt_tga)
            except Exception as e:
                st.error(f"Failed to generate analysis: {e}")
                result = ""
            st.session_state['target_group_analysis'] = result.strip()
        if st.session_state.get('target_group_analysis'):
            st.text_area("Target Group Analysis", value=st.session_state['target_group_analysis'], height=200)
            # Download buttons
            txt_data = st.session_state['target_group_analysis']
            st.download_button("Download as TXT", data=txt_data, file_name="target_group_analysis.txt")
            pdf_bytes = text_to_pdf_bytes(txt_data)
            if pdf_bytes:
                st.download_button("Download as PDF", data=pdf_bytes, file_name="target_group_analysis.pdf", mime="application/pdf")
    with tab2:
        if st.button("Generate Job Advertisement"):
            # Prepare prompt for job ad
            jt = st.session_state['job_title']; cn = st.session_state['company_name']; loc = st.session_state['city']
            rd = st.session_state['role_description']; kr = st.session_state['key_responsibilities']
            req_skills = st.session_state['must_have_skills'] or st.session_state['hard_skills']
            tone = st.session_state['ad_seniority_tone']; length = st.session_state['ad_length_preference']
            # Summarize benefits from state
            benefits = []
            if st.session_state['flexible_hours'] in ["Yes", "Partial/Flex Schedule"]: benefits.append("flexible working hours")
            if st.session_state['remote_work_policy'] in ["Hybrid", "Full Remote"]: benefits.append(f"{st.session_state['remote_work_policy']} work options")
            if st.session_state['relocation_assistance'] == "Yes": benefits.append("relocation assistance")
            if st.session_state['childcare_support'] == "Yes": benefits.append("childcare support")
            if st.session_state['vacation_days']: benefits.append(f"{st.session_state['vacation_days']} paid vacation days")
            benefits_str = ", ".join(benefits)
            prompt_ad = (
                f"Write a job advertisement for the following position:\n"
                f"Job Title: {jt}\n" + (f"Company: {cn}\n" if cn else "") + (f"Location: {loc}\n" if loc else "") +
                (f"Role Description: {rd}\n" if rd else "") + (f"Key Responsibilities: {kr}\n" if kr else "") +
                (f"Required Skills: {req_skills}\n" if req_skills else "") +
                (f"Benefits: {benefits_str}\n" if benefits_str else "") +
                f"Tone: {tone}. Length: {length}.\n"
                "The ad should be engaging and include a brief company intro, role responsibilities, required qualifications, any benefits, and a call to action.")
            try:
                result = llm.complete(prompt=prompt_ad)
            except Exception as e:
                st.error(f"Failed to generate job ad: {e}")
                result = ""
            st.session_state['generated_job_ad'] = result.strip()
        if st.session_state.get('generated_job_ad'):
            st.text_area("Job Advertisement", value=st.session_state['generated_job_ad'], height=250)
            txt_data = st.session_state['generated_job_ad']
            st.download_button("Download as TXT", data=txt_data, file_name="job_ad.txt")
            pdf_bytes = text_to_pdf_bytes(txt_data)
            if pdf_bytes:
                st.download_button("Download as PDF", data=pdf_bytes, file_name="job_ad.pdf", mime="application/pdf")
    with tab3:
        if st.button("Generate Interview Prep Guide"):
            # Prepare prompt for interview prep
            jt = st.session_state['job_title']; kr = st.session_state['key_responsibilities']
            must_skills = st.session_state['must_have_skills']; soft_sk = st.session_state['soft_skills']
            prompt_prep = (
                f"You are preparing to interview candidates for the position of {jt}. "
                "Based on the job description, create an interview preparation guide for the interviewer. "
                "Include:\n1. A brief overview of what to look for in a candidate.\n"
                "2. 5-10 key interview questions (technical and behavioral) to assess required skills and competencies.\n"
                "3. For each question, notes on what a good answer should include.\n"
            )
            # Provide some context from the spec if available
            if kr: prompt_prep += f"Key Responsibilities: {kr}\n"
            if must_skills: prompt_prep += f"Must-Have Skills: {must_skills}\n"
            if soft_sk: prompt_prep += f"Desired Soft Skills: {soft_sk}\n"
            try:
                result = llm.complete(prompt=prompt_prep)
            except Exception as e:
                st.error(f"Failed to generate interview prep: {e}")
                result = ""
            st.session_state['generated_interview_prep'] = result.strip()
        if st.session_state.get('generated_interview_prep'):
            st.text_area("Interview Preparation Guide", value=st.session_state['generated_interview_prep'], height=300)
            txt_data = st.session_state['generated_interview_prep']
            st.download_button("Download as TXT", data=txt_data, file_name="interview_prep.txt")
            pdf_bytes = text_to_pdf_bytes(txt_data)
            if pdf_bytes:
                st.download_button("Download as PDF", data=pdf_bytes, file_name="interview_prep.pdf", mime="application/pdf")
    # Export session data as JSON
    export_data = {k: v for k, v in st.session_state.items() if k != "uploaded_file"}
    st.download_button("Download All Data (JSON)", data=json.dumps(export_data, indent=2), file_name="vacalyser_session.json", mime="application/json")
    # Additional Tools
    st.subheader("Additional Tools (Optional)")
    # Responsibility Triangle Visualization
    tech_count = len(st.session_state['technical_tasks'].split())
    mgr_count = len(st.session_state['managerial_tasks'].split())
    admin_count = len(st.session_state['administrative_tasks'].split())
    total = tech_count + mgr_count + admin_count
    if total > 0:
        a = tech_count / total if total else 0
        b = mgr_count / total if total else 0
        c = admin_count / total if total else 0
        fig = px.scatter_ternary(a=[a], b=[b], c=[c])
        fig.update_traces(marker=dict(size=15, color="blue"))
        fig.update_layout(ternary=dict(aaxis_title="Technical", baxis_title="Managerial", caxis_title="Administrative"))
        st.write("**Responsibility Distribution (Technical vs Managerial vs Administrative)**")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.write("*Responsibility distribution will display here once task details are provided.*")
    # Email Template Generator
    with st.expander("Email Template Generator"):
        if st.button("Generate Outreach Email"):
            jt = st.session_state['job_title']; cn = st.session_state['company_name']
            prompt_email = (
                f"Write a concise, professional recruitment email to a potential candidate for the position of {jt} at {cn if cn else 'our company'}. "
                "Introduce the company briefly, highlight the role's key points (responsibilities or benefits), and include a friendly call to action to apply."
            )
            try:
                result = llm.complete(prompt=prompt_email)
            except Exception as e:
                st.error(f"Failed to generate email: {e}")
                result = ""
            st.session_state['generated_email_template'] = result.strip()
        if st.session_state.get('generated_email_template'):
            st.text_area("Email Template", value=st.session_state['generated_email_template'], height=150)
            st.download_button("Download Email as TXT", data=st.session_state['generated_email_template'], file_name="email_template.txt")
    # Boolean Search Generator
    with st.expander("Boolean Search Query Generator"):
        if st.button("Generate Boolean Search Query"):
            jt = st.session_state['job_title']; must_skills = st.session_state['must_have_skills']
            prompt_bool = (
                f"Generate a Boolean search string to find resumes for a {jt} role "
                + (f"requiring skills: {must_skills}. " if must_skills else "")
                + "Use AND, OR, and NOT operators appropriately."
            )
            try:
                result = llm.complete(prompt=prompt_bool)
            except Exception as e:
                st.error(f"Failed to generate search query: {e}")
                result = ""
            st.session_state['generated_boolean_query'] = result.strip()
        if st.session_state.get('generated_boolean_query'):
            st.code(st.session_state['generated_boolean_query'], language="")

def main():
    st.set_page_config(page_title="Vacalyser Wizard", layout="wide")
    apply_base_styling()
    show_sidebar_links()
    if "wizard_step" not in st.session_state:
        initialize_session_state()
    step = st.session_state.get("wizard_step", 1)
    if step == 1:
        start_discovery_page()
    elif step == 2:
        render_step_2()
    elif step == 3:
        render_step_3()
    elif step == 4:
        render_step_4()
    elif step == 5:
        render_step_5()
    elif step == 6:
        render_step_6()
    elif step == 7:
        render_step_7()
    elif step == 8:
        render_step_8()
    # Navigation buttons
    if step > 1:
        if st.button("⬅️ Back"):
            st.session_state["wizard_step"] -= 1
            st.experimental_rerun()
    if step < 8:
        if st.button("Next ➡"):
            st.session_state["wizard_step"] += 1
            st.experimental_rerun()

if __name__ == "__main__":
    main()
